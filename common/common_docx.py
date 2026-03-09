# -*- coding: utf-8 -*-
'''
@文件: common_docx.py
@說明:
@時間: 2024/11/06 15:41:47
@作者: LiDong
'''

import os
from io import BytesIO

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from common.common_minio import OperMinio
from configs.constant import BUCKET


class AdvancedWordTableGenerator:

    def __init__(self, file_path=None):
        if file_path and os.path.exists(file_path):
            self.document = Document(file_path)
        else:
            self.document = Document()
        self.om = OperMinio()

    def add_title(self, title: str) -> None:
        title_paragraph = self.document.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(13)
        title_run.font.name = 'Times New Roman'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

    def add_headers(self, table, headers):
        for i, row_data in enumerate(headers, start=0):
            row_cells = table.rows[i].cells
            for _index, data in enumerate(row_data):
                paragraph = row_cells[_index].paragraphs[0]
                run = paragraph.add_run(str(data))
                run.font.size = Pt(10.5)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                run.bold = True

    def create_table(self, rows, cols=2, column_widths=[4, 8]):
        table = self.document.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        for i, width in enumerate(column_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)
        return table

    def __set_content_font(self, obj, color=None, bold=False):
        obj.font.bold = bold
        obj.font.size = Pt(10.5)
        obj.font.name = 'Times New Roman'
        obj._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        if color is not None:
            obj.font.color.rgb = RGBColor(color[0], color[1], color[2])

    def _add_cell_content(self, cell, content, bold=False, color=None):
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(content)
        self.__set_content_font(run, color, bold)

    def __add_function_data(self, function_list, right_cell):
        content_paragraph = right_cell.paragraphs[0]
        for func in function_list:
            obj = content_paragraph.add_run(f"{func['function_nm']} {func['function_progress']} {func['time_consum']}h\n")
            self.__set_content_font(obj)
            obj = content_paragraph.add_run(f"- {func['progress_record']}\n")
            self.__set_content_font(obj)

            # 添加图片（从MinIO获取）
            for path in func.get("_paths", list()):
                # 从MinIO获取图片数据
                if path.startswith("/") or path.startswith("\\"):
                    path = path[1:]
                filelist = self.om.search_files(BUCKET, path + "/files")
                for file in filelist:
                    if file.split(".")[-1].lower() not in ["jpg", "png", "jpeg", "svg"]:
                        continue
                    response = self.om.use_filename_get_stream(BUCKET, file)
                    if not response:
                        continue
                    image_data = BytesIO(response)
                    # 添加到文档
                    run = content_paragraph.add_run()
                    run.add_picture(image_data, width=Inches(3.8))
                    content_paragraph.add_run("\n") # 空行

    def save(self, file_path=None):
        self.document.save(file_path)

    def run(self, headers, title, datalist):
        self.add_title(title)
        start_rows = len(headers)
        rows = start_rows + len(datalist)
        table = self.create_table(rows)
        self.add_headers(table, headers)
        for i, row_data in enumerate(datalist, start=start_rows):
            row_cells = table.rows[i].cells
            left_cell = row_cells[0]
            content = f"{i + 1 - start_rows}. {row_data[0]}"
            self._add_cell_content(left_cell, content, bold=True)
            right_cell = row_cells[1]
            self.__add_function_data(row_data[1], right_cell)

        return table
